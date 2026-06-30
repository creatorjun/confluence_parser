# to_docx.py
from __future__ import annotations

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

from config import ROOT_PAGE_ID
from confluence_client import collect_pages
from html_cleaner import clean, is_code_table, extract_code_lines

OUTPUT_FILE = "output.docx"


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _process_inline(para, node):
    if isinstance(node, NavigableString):
        text = str(node)
        if text.strip():
            para.add_run(text)
        return
    tag = node.name
    if tag in ("strong", "b"):
        run = para.add_run(node.get_text())
        run.bold = True
    elif tag in ("em", "i"):
        run = para.add_run(node.get_text())
        run.italic = True
    elif tag == "code":
        run = para.add_run(node.get_text())
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    elif tag == "a":
        run = para.add_run(node.get_text())
        run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
    elif tag == "br":
        para.add_run("\n")
    else:
        for child in node.children:
            _process_inline(para, child)


def _add_code_block(doc: Document, text: str):
    text = text.strip()
    if not text:
        return
    lines = text.split("\n")
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "4A90D9")
    pBdr.append(left)
    pPr.append(pBdr)
    for i, line in enumerate(lines):
        run = para.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if i < len(lines) - 1:
            br = OxmlElement("w:br")
            run._r.append(br)


def _add_table(doc: Document, table_tag: Tag):
    rows = table_tag.find_all("tr")
    if not rows:
        return
    max_cols = max(
        sum(int(td.get("colspan", 1)) for td in row.find_all(["td", "th"]))
        for row in rows
    )
    if max_cols == 0:
        return
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        c_idx = 0
        for cell_tag in cells:
            if c_idx >= max_cols:
                break
            cell = table.cell(r_idx, c_idx)
            is_header = cell_tag.name == "th"
            if is_header:
                _set_cell_bg(cell, "E8F0FE")
            _set_cell_border(cell)
            para = cell.paragraphs[0]
            para.clear()
            for child in cell_tag.children:
                _process_inline(para, child)
            for run in para.runs:
                run.font.size = Pt(9)
                if is_header:
                    run.bold = True
            c_idx += int(cell_tag.get("colspan", 1))


def _process_block(doc: Document, node):
    if isinstance(node, NavigableString):
        return
    tag = node.name
    if tag is None:
        return

    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(tag[1])
        heading_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3",
                       4: "Heading 4", 5: "Heading 5", 6: "Heading 6"}
        para = doc.add_paragraph(style=heading_map.get(level, "Heading 3"))
        para.clear()
        for child in node.children:
            _process_inline(para, child)
    elif tag == "p":
        if not node.get_text(strip=True):
            return
        para = doc.add_paragraph()
        for child in node.children:
            _process_inline(para, child)
    elif tag in ("pre", "code"):
        _add_code_block(doc, node.get_text())
    elif tag == "div":
        classes = node.get("class", [])
        if any(c in classes for c in ("code", "codeContent", "syntaxhighlighter", "highlight")):
            code_node = node.find("pre") or node.find("code")
            _add_code_block(doc, code_node.get_text() if code_node else node.get_text())
        else:
            for child in node.children:
                _process_block(doc, child)
    elif tag == "table":
        if is_code_table(node):
            _add_code_block(doc, extract_code_lines(node))
        else:
            _add_table(doc, node)
        doc.add_paragraph()
    elif tag in ("ul", "ol"):
        for li in node.find_all("li", recursive=False):
            para = doc.add_paragraph(
                style="List Bullet" if tag == "ul" else "List Number"
            )
            for child in li.children:
                _process_inline(para, child)
    elif tag == "hr":
        doc.add_paragraph("─" * 60)
    else:
        for child in node.children:
            _process_block(doc, child)


def build_docx(pages: list[tuple[int, str, str]]) -> Document:
    doc = Document()
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10)
    heading_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
    for depth, title, html_body in pages:
        level = min(depth + 1, 4)
        doc.add_paragraph(title, style=heading_map[level])
        soup = BeautifulSoup(html_body, "html.parser")
        soup = clean(soup)
        for child in soup.children:
            _process_block(doc, child)
        if depth == 0:
            doc.add_page_break()
    return doc


def main():
    if not ROOT_PAGE_ID:
        print("[ERROR] CONFLUENCE_ROOT_PAGE_ID 환경변수를 설정하세요.")
        return
    print(f"페이지 수집 중... (root={ROOT_PAGE_ID})\n")
    pages = collect_pages(ROOT_PAGE_ID)
    print(f"\n총 {len(pages)}개 페이지 수집. docx 생성 중...")
    doc = build_docx(pages)
    doc.save(OUTPUT_FILE)
    print(f"완료: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
