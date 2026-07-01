# infrastructure/converters/docx_converter.py
from __future__ import annotations

from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from domain.model import PageNode
from domain.ports import IConverter
from infrastructure.converters.html_cleaner import clean, is_code_table, extract_code_lines


def _shd(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), hex_color)
    tcPr.append(s)


def _border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    b = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "AAAAAA")
        b.append(el)
    tcPr.append(b)


def _inline(para, node):
    if isinstance(node, NavigableString):
        t = str(node)
        if t.strip():
            para.add_run(t)
        return
    tag = node.name
    if tag in ("strong", "b"):
        para.add_run(node.get_text()).bold = True
    elif tag in ("em", "i"):
        r = para.add_run(node.get_text())
        r.italic = True
    elif tag == "code":
        r = para.add_run(node.get_text())
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    elif tag == "a":
        r = para.add_run(node.get_text())
        r.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
    elif tag == "br":
        para.add_run("\n")
    else:
        for c in node.children:
            _inline(para, c)


def _code_block(doc: Document, text: str):
    text = text.strip()
    if not text:
        return
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Inches(0.3)
    pPr = para._p.get_or_add_pPr()
    s = OxmlElement("w:shd")
    s.set(qn("w:val"), "clear")
    s.set(qn("w:color"), "auto")
    s.set(qn("w:fill"), "F4F4F4")
    pPr.append(s)
    pb = OxmlElement("w:pBdr")
    lft = OxmlElement("w:left")
    lft.set(qn("w:val"), "single")
    lft.set(qn("w:sz"), "12")
    lft.set(qn("w:space"), "4")
    lft.set(qn("w:color"), "4A90D9")
    pb.append(lft)
    pPr.append(pb)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        r = para.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        if i < len(lines) - 1:
            r._r.append(OxmlElement("w:br"))


def _table(doc: Document, tbl: Tag):
    rows = tbl.find_all("tr")
    if not rows:
        return
    max_cols = max(
        sum(int(td.get("colspan", 1)) for td in row.find_all(["td", "th"]))
        for row in rows
    )
    if not max_cols:
        return
    t = doc.add_table(rows=len(rows), cols=max_cols)
    t.style = "Table Grid"
    for ri, row in enumerate(rows):
        ci = 0
        for ct in row.find_all(["td", "th"]):
            if ci >= max_cols:
                break
            cell = t.cell(ri, ci)
            if ct.name == "th":
                _shd(cell, "E8F0FE")
            _border(cell)
            para = cell.paragraphs[0]
            para.clear()
            for ch in ct.children:
                _inline(para, ch)
            for r in para.runs:
                r.font.size = Pt(9)
                if ct.name == "th":
                    r.bold = True
            ci += int(ct.get("colspan", 1))


def _block(doc: Document, node):
    if isinstance(node, NavigableString):
        return
    tag = node.name
    if not tag:
        return
    _HM = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4", 5: "Heading 5", 6: "Heading 6"}
    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        p = doc.add_paragraph(style=_HM[int(tag[1])])
        p.clear()
        for c in node.children:
            _inline(p, c)
    elif tag == "p":
        if not node.get_text(strip=True):
            return
        p = doc.add_paragraph()
        for c in node.children:
            _inline(p, c)
    elif tag in ("pre", "code"):
        _code_block(doc, node.get_text())
    elif tag == "div":
        classes = node.get("class", [])
        if any(c in classes for c in ("code", "codeContent", "syntaxhighlighter", "highlight")):
            cn = node.find("pre") or node.find("code")
            _code_block(doc, cn.get_text() if cn else node.get_text())
        else:
            for c in node.children:
                _block(doc, c)
    elif tag == "table":
        if is_code_table(node):
            _code_block(doc, extract_code_lines(node))
        else:
            _table(doc, node)
        doc.add_paragraph()
    elif tag in ("ul", "ol"):
        style = "List Bullet" if tag == "ul" else "List Number"
        for li in node.find_all("li", recursive=False):
            p = doc.add_paragraph(style=style)
            for c in li.children:
                _inline(p, c)
    elif tag == "hr":
        doc.add_paragraph("─" * 60)
    else:
        for c in node.children:
            _block(doc, c)


class DocxConverter(IConverter):
    def convert(self, pages: list[PageNode], output_path: str) -> None:
        doc = Document()
        doc.styles["Normal"].font.name = "맑은 고딕"
        doc.styles["Normal"].font.size = Pt(10)
        HM = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
        for page in pages:
            doc.add_paragraph(page.title, style=HM[min(page.depth + 1, 4)])
            soup = clean(BeautifulSoup(page.html_body, "html.parser"))
            for c in soup.children:
                _block(doc, c)
            if page.depth == 0:
                doc.add_page_break()
        doc.save(output_path)
